from docx import Document
from docx.shared import RGBColor


def main(names, file_address):
    def set_run(run):
        # 设置run的字体大小、是否加粗以及字体颜色
        run.font.size = font_size
        run.bold = bold
        run.font.color.rgb = color
    # 读取Word文档
    file_address = './task2_data/' + file_address
    doc = Document(file_address)
    # names = ['电力大学', '银行']
    for i in range(len(names)):
        for p in doc.paragraphs:
            for r in p.runs:
                if names[i] not in r.text:
                    pass
                # 获取当前run的字体属性
                font_size = r.font.size
                bold = r.bold
                color = r.font.color.rgb
                # 使用关键词切分当前run的文本

                rest = r.text.split(names[i])
                #清除当前run的内容
                r.text = ''
                for text in rest[:-1]:
                    run = p.add_run(text=text)
                    set_run(run)
                    run = p.add_run(names[i])
                    run.font.size = font_size
                    run.bold = bold
                    run.font.color.rgb = RGBColor(255, 0, 0)
                run = p.add_run(rest[-1])
                set_run(run)
    global save_address
    save_address = './task2_data/handle.docx'
    doc.save(save_address)

def readword():
    return save_address

