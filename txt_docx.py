from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
import os

document = Document()

document.styles['Normal'].font.name = '宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
document.styles['Normal'].font.size = Pt(10.5)
document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)


paragraph = document.add_paragraph()
run = paragraph.add_run()
#run.font.name = 'Times New Roman'
run.font.name = 'Cambria'
run.font.color.rgb = RGBColor(0,0,0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')

with open("./task2_data/test.txt", "r", encoding='UTF-8') as f:
    line = f.readline()
    while line:
        document.add_paragraph(line)
        line = f.readline()
print('任务完成！！')

# 保存文档
document.save('text.docx')

