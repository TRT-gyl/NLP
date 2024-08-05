import string
import nltk.data
from docx import Document
from CorrectWords import correct_text_generic
from docx.shared import RGBColor


punkt_list = r",.?\"'!()/\\-<>:@#$%^&*~"
document = Document()  # word文档句柄
count_correct = 0


def strip_punctuation(text):
    return ''.join([c for c in text if c not in string.punctuation])


def tokenizes(paragraph):
    tokenizer = nltk.data.load('./task1_data/english.pickle')
    list_ret = list()
    for row in tokenizer.tokenize(paragraph):
        list_ret.append(row)
    return(list_ret)


def words(sentences):
    word_list = []
    for i in range(len(sentences)):
        sentence = strip_punctuation(sentences[i])
        word_list.append(sentence.split())
    return(word_list)


def write_correct_paragraph(file_address):
    # global file_address
    file_address = './task1_data/' + file_address
    file = Document(file_address)
    # print("段落数：" + str(len(file.paragraphs)))
    global count_correct
    # 文档中修改的单词个数
    count_correct = 0
    for i in range(len(file.paragraphs)):
        # 每一段的内容
        paragraph = file.paragraphs[i].text.strip()
        # 进行句子划分
        sentences = tokenizes(paragraph)
        # 词语划分
        words_list = words(sentences)
        # 段落句柄
        p = document.add_paragraph(' ' * 7)
        for word_list in words_list:
            for word in word_list:
                if word not in punkt_list:

                    p.add_run(' ')
                    # 纠正单词，如果单词正确，则返回原单词
                    correct_word = correct_text_generic(word)

                    # 每一句话第一个单词的第一个字母大写
                    if word_list.index(word) == 0 and words_list.index(word_list) == 0:
                        correct_word = correct_word[0].upper() + correct_word[1:]

                    # 如果单词有修改，则颜色为红色
                    if correct_word != word:
                        colored_word = p.add_run(correct_word)
                        font = colored_word.font
                        font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        count_correct += 1
                    else:
                        p.add_run(correct_word)
                else:
                    p.add_run(word)
    global save_address
    save_address = './task1_data/CorrectDocument.docx'
    document.save(save_address)
    print("一共修改了%d处错误单词。" % count_correct)
    print("修改并保存文件完毕！" + save_address)


def get_count_correct():
    return count_correct


def readword():
    return save_address


if __name__ == '__main__':

    write_correct_paragraph('ErrorDocument.docx')
    print("修改并保存文件完毕！")
    print("一共修改了%d处。" % count_correct)